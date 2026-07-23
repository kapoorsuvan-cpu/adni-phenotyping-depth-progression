#!/usr/bin/env python3
"""Build the journal manuscript and supplementary material from final outputs."""

from __future__ import annotations

from datetime import date
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT = Path(__file__).resolve().parents[1]
OUTPUTS = PROJECT / "final_report_outputs"
TABLES = OUTPUTS / "tables"
FIGURES = OUTPUTS / "figures"
PUBLICATION = PROJECT / "publication"
MANUSCRIPT = PUBLICATION / "Kapoor_et_al_ADNI_progression_manuscript.docx"
SUPPLEMENT = PUBLICATION / "Kapoor_et_al_ADNI_progression_supplement.docx"

NAVY = "17324D"
BLUE = "2D5B7B"
LIGHT_BLUE = "EAF1F6"
LIGHT_GRAY = "F3F5F7"
MID_GRAY = "D7DDE2"
TEXT = RGBColor(34, 43, 51)
MUTED = RGBColor(82, 96, 108)


def load_table(name: str) -> pd.DataFrame:
    return pd.read_csv(TABLES / name)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=95, bottom=80, end=95) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_table_borders(table, color=MID_GRAY, size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def configure_document(doc: Document, short_title: str) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.30)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.3)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(5.5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for name, size, color, before, after in (
        ("Title", 20, NAVY, 0, 9),
        ("Heading 1", 14, NAVY, 14, 5),
        ("Heading 2", 11.5, BLUE, 9, 3),
        ("Heading 3", 10.5, BLUE, 7, 2),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("Caption", "Quote"):
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(9)
        style.font.color.rgb = MUTED
        style.paragraph_format.space_after = Pt(7)

    if "Table note" not in styles:
        table_note = styles.add_style("Table note", WD_STYLE_TYPE.PARAGRAPH)
        table_note.font.name = "Arial"
        table_note.font.size = Pt(8.4)
        table_note.font.color.rgb = MUTED
        table_note.paragraph_format.space_after = Pt(6)

    header = section.header.paragraphs[0]
    header.text = short_title
    header.style = styles["Caption"]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.runs[0].font.color.rgb = RGBColor.from_string(BLUE)
    add_page_number(section.footer.paragraphs[0])

    core = doc.core_properties
    core.title = short_title
    core.subject = "Prediction of 24-month progression from mild cognitive impairment to dementia"
    core.author = "Suvan Kapoor; Dominic Ablakhad; Rayan Hanna; Kylan Huynh; Eric Quirarte; Orion Nocon"
    core.keywords = "mild cognitive impairment, dementia, ADNI, prediction, MRI, APOE, biomarkers"


def add_label_paragraph(doc: Document, label: str, text: str, style=None) -> None:
    p = doc.add_paragraph(style=style)
    r = p.add_run(label)
    r.bold = True
    p.add_run(text)


def add_rule(doc: Document, color=BLUE) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    p_pr.append(pbdr)


def add_table(
    doc: Document,
    title: str,
    headers: list[str],
    rows: list[list[str]],
    widths: list[int],
    note: str,
    font_size=8.3,
) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(NAVY)

    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_borders(table)
    header = table.rows[0]
    set_repeat_table_header(header)
    prevent_row_split(header)
    for idx, (cell, label) in enumerate(zip(header.cells, headers)):
        set_cell_width(cell, widths[idx])
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(label))
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(font_size)

    for row_idx, values in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        if row_idx % 2:
            for cell in row.cells:
                set_cell_shading(cell, LIGHT_GRAY)
        for idx, (cell, value) in enumerate(zip(row.cells, values)):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            r.font.size = Pt(font_size)
    p = doc.add_paragraph(note, style="Table note")
    p.paragraph_format.keep_together = True


def add_figure(doc: Document, number: int, filename: str, caption: str, width=6.8) -> None:
    image = FIGURES / filename
    if not image.exists():
        raise FileNotFoundError(image)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(str(image), width=Inches(width))
    cap = doc.add_paragraph(style="Caption")
    cap.paragraph_format.keep_together = True
    cap.paragraph_format.space_after = Pt(9)
    r = cap.add_run(f"Figure {number}. ")
    r.bold = True
    cap.add_run(caption)


def fmt_ci(value, low, high, digits=3) -> str:
    return f"{value:.{digits}f} ({low:.{digits}f}–{high:.{digits}f})"


def fmt_q(value) -> str:
    if pd.isna(value):
        return "—"
    if value < 0.001:
        return "<0.001"
    return f"{value:.3f}"


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("ORIGINAL RESEARCH")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(BLUE)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(
        "How Much Phenotyping Is Needed to Predict 24-Month Progression "
        "From Mild Cognitive Impairment to Dementia?"
    )

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(13)
    run = subtitle.add_run(
        "A leakage-resistant ADNI analysis with stepwise multimodal ablation "
        "and OASIS-2 conceptual replication"
    )
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED

    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors.paragraph_format.space_after = Pt(8)
    arun = authors.add_run(
        "Suvan Kapoor; Dominic Ablakhad; Rayan Hanna; Kylan Huynh; "
        "Eric Quirarte; Orion Nocon; for the Alzheimer’s Disease "
        "Neuroimaging Initiative*"
    )
    arun.bold = True
    arun.font.size = Pt(11)

    affiliation = doc.add_paragraph()
    affiliation.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affiliation.add_run("Author affiliations: ").bold = True
    affiliation.add_run("to be confirmed before journal submission.")
    corresponding = doc.add_paragraph()
    corresponding.alignment = WD_ALIGN_PARAGRAPH.CENTER
    corresponding.add_run("Corresponding author: ").bold = True
    corresponding.add_run("Suvan Kapoor; postal address and email to be supplied.")

    add_rule(doc)
    add_label_paragraph(doc, "Word count: ", "approximately 4,600 words, excluding references, tables, and legends.")
    add_label_paragraph(doc, "Tables and figures: ", "4 main tables, 7 figures, and supplementary tables.")
    add_label_paragraph(doc, "Short title: ", "Phenotyping depth in MCI progression")
    add_label_paragraph(doc, "Keywords: ", "mild cognitive impairment; dementia; ADNI; prediction model; cognition; MRI; APOE; amyloid; p-tau")
    add_label_paragraph(doc, "Reporting framework: ", "TRIPOD+AI; PROBAST+AI principles.")

    doc.add_page_break()


def build_main() -> None:
    t1 = load_table("table1_baseline_characteristics.csv")
    t3 = load_table("table3_stepwise_classification.csv")
    t4 = load_table("table4_continuous_outcomes.csv")
    t5 = load_table("table5_minimal_marker_panels.csv")
    t6 = load_table("table6_outcome_descriptives.csv")
    t7 = load_table("table7_sensitivity_cohorts.csv")
    t8 = load_table("table8_oasis2_external_replication.csv")
    t9 = load_table("table9_subgroup_performance.csv")

    doc = Document()
    configure_document(doc, "Phenotyping depth in MCI progression")
    add_title_page(doc)

    doc.add_heading("Abstract", level=1)
    add_label_paragraph(
        doc,
        "Background: ",
        "Prediction studies often combine clinical tests, magnetic resonance imaging (MRI), "
        "genotype, and Alzheimer disease (AD) biomarkers without testing how much each layer "
        "adds. We quantified the incremental value of baseline phenotyping for 24-month "
        "progression from mild cognitive impairment (MCI) to all-cause dementia.",
    )
    add_label_paragraph(
        doc,
        "Methods: ",
        "We formed a strict longitudinal cohort from the Alzheimer’s Disease Neuroimaging "
        "Initiative (ADNI). The index was each participant’s first MCI visit. Progressors had "
        "a dementia diagnosis within 24 months. Stable participants required at least 24 "
        "months of follow-up without dementia. Nine cumulative feature sets were assessed "
        "with regularized logistic regression. Preprocessing and hyperparameter selection "
        "occurred within nested cross-validation. Adjacent changes in area under the receiver "
        "operating characteristic curve (AUROC) were tested with paired stratified bootstrap "
        "samples and Benjamini–Hochberg correction. A compact panel used a locked 25% test "
        "set. Continuous MMSE, MoCA, and CDR Sum of Boxes (CDR-SB) changes were secondary "
        "outcomes. OASIS-2 supplied a conceptual replication.",
    )
    add_label_paragraph(
        doc,
        "Results: ",
        "The ADNI cohort included 1,004 participants; 261 progressed to dementia. "
        "Demographics alone yielded AUROC 0.510. Adding MMSE increased AUROC to 0.675 "
        "(increment 0.178; false-discovery-rate-adjusted q<0.001). MoCA, CDR-SB, and the "
        "broader cognitive battery produced further increments. Broad cognition reached "
        "AUROC 0.866. MRI changed AUROC by −0.000 (q=0.970). APOE changed AUROC by 0.005 "
        "(q=0.091). AD biomarkers changed AUROC by 0.005 (q=0.255). The full model yielded "
        "AUROC 0.875 (95% bootstrap confidence interval 0.852–0.897), precision–recall area "
        "0.717, and Brier score 0.123. A five-marker clinical panel yielded locked-test AUROC "
        "0.853. Its multimodal counterpart yielded 0.844 (difference −0.008; q=1.000). "
        "OASIS-2 included 34 eligible participants and did not provide a precise replication "
        "(full AUROC 0.637; 95% confidence interval 0.443–0.817).",
    )
    add_label_paragraph(
        doc,
        "Conclusions: ",
        "Cognitive and functional measures carried most of the short-horizon predictive "
        "information in this ADNI cohort. Available MRI, APOE, and biomarker variables did "
        "not improve discrimination after correction for multiple testing. This is a "
        "conditional prediction result, not evidence that these measures lack biological "
        "or diagnostic value. External validation in a larger biomarker-rich cohort remains "
        "required.",
    )

    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(
        "Mild cognitive impairment is a clinical syndrome with variable causes and courses. "
        "Some people remain stable. Others progress to dementia over a short interval. The "
        "transition has direct relevance to counseling, trial recruitment, and follow-up "
        "planning [1–3]. Clinical stage and biological AD stage are related but distinct. "
        "Current criteria define AD through biomarkers and recognize that symptoms are "
        "modified by reserve and copathology [14]. A prediction target based on clinical "
        "conversion should therefore be described as progression to dementia rather than "
        "conversion to biological AD."
    )
    doc.add_paragraph(
        "Published models use demographics, cognition, structural MRI, APOE, cerebrospinal "
        "fluid, positron emission tomography, and blood biomarkers. Reported discrimination "
        "varies, and external validation is uncommon [9–13]. Many studies evaluate a final "
        "multimodal model without asking whether added testing changes prediction once "
        "routine clinical measures are known. That question affects cost, participant burden, "
        "and feasibility. It also separates biological association from incremental "
        "prediction."
    )
    doc.add_paragraph(
        "Oxtoby and colleagues used a data-driven disease progression model to stage trial "
        "participants and study targeted screening [9]. Van der Veere and colleagues modeled "
        "MMSE trajectories in amyloid-positive MCI and mild dementia, then tested transport "
        "to ADNI [10]. Their work shows two needs. The outcome must be defined before model "
        "development. The contribution of each phenotyping layer must be measured under "
        "validation."
    )
    doc.add_paragraph(
        "We addressed these needs with a prespecified stepwise question: how much baseline "
        "phenotyping is needed to predict all-cause dementia within 24 months after the first "
        "MCI visit? We expected cognitive and functional tests to provide the first change in "
        "performance. We then tested whether MRI, APOE, and available AD biomarkers added "
        "discrimination. Secondary analyses assessed continuous change, calibration, decision "
        "utility, compact panels, subgroup performance, and conceptual replication in OASIS-2."
    )

    doc.add_heading("Methods", level=1)
    doc.add_heading("Study design and data sources", level=2)
    doc.add_paragraph(
        "This retrospective prediction study used de-identified longitudinal ADNI data. Data "
        "used in the preparation of this article were obtained from the Alzheimer’s Disease "
        "Neuroimaging Initiative database (adni.loni.usc.edu). ADNI began in 2003 as a "
        "public–private partnership led by Michael W. Weiner, MD. Its goal has been to test "
        "whether serial MRI, positron emission tomography, other biological markers, and "
        "clinical and neuropsychological assessments can be combined to measure progression "
        "of MCI and early AD [3]."
    )
    doc.add_paragraph(
        "OASIS-2 supplied the external conceptual replication. It contains longitudinal MRI "
        "and clinical data from 150 adults aged 60–96 years [16]. The official Washington "
        "University subject-data workbook was downloaded by the pipeline. A SHA-256 checksum "
        "was recorded. OASIS-2 lacks the ADNI cognitive battery and AD biomarker set, so it "
        "cannot provide direct transport validation."
    )

    doc.add_heading("Cohort and outcomes", level=2)
    doc.add_paragraph(
        "The index date was the first visit classified as MCI. Visit codes were converted to "
        "elapsed month. Screening, baseline, screening MRI, init, m0, and m00 were treated as "
        "month 0. All visit-level features were joined on participant identifier and "
        "harmonized visit month. A progressor had a dementia diagnosis after the index visit "
        "and no later than month 24. A stable participant had no dementia diagnosis and at "
        "least 24 months of observed follow-up. Participants without follow-up were excluded. "
        "Participants with less than 24 months of follow-up and no observed progression were "
        "also excluded."
    )
    doc.add_paragraph(
        "Secondary outcomes were raw score change from baseline to approximately 24 months. "
        "Baseline scores were selected within ±3 months of the MCI index. Follow-up scores "
        "were selected between 18 and 30 months, with the visit closest to month 24 retained. "
        "No later score replaced a missing baseline. Outcomes were MMSE, MoCA, and CDR-SB "
        "[4–6]."
    )

    doc.add_heading("Predictor ladder", level=2)
    doc.add_paragraph(
        "Nine cumulative feature sets represented increasing phenotyping depth: "
        "(1) age, sex, and education; (2) depressive symptoms; (3) MMSE; (4) MoCA; "
        "(5) CDR-SB; (6) a broad cognitive and functional battery; (7) regional structural "
        "MRI measures; (8) APOE ε4 allele count; and (9) AD biomarkers. The broad battery "
        "included FAQ, ADAS-Cog measures, memory tests, language measures, executive tests, "
        "and ADNI composite scores [7,8]. The final biomarker layer contained amyloid PET "
        "status and Centiloids, plasma p-tau181, legacy and Roche CSF amyloid and tau assays, "
        "and derived ratios. Plasma p-tau217 had eight usable baseline records and was "
        "excluded before modeling."
    )

    doc.add_heading("Model development and validation", level=2)
    doc.add_paragraph(
        "The classifier was L2-regularized logistic regression. This model was chosen to "
        "estimate the contribution of each measurement layer with limited algorithmic "
        "complexity. Missing numeric values were median-imputed and accompanied by missingness "
        "indicators. Variables were standardized. These operations were fit within each "
        "training fold."
    )
    doc.add_paragraph(
        "The primary estimates came from five outer stratified folds. Each outer training "
        "set contained a four-fold grid search over inverse regularization strengths "
        "C={0.003, 0.01, 0.03, 0.1, 0.3, 1}. Negative log loss selected C. Predictions for "
        "each participant were produced only when that participant was in an outer test fold. "
        "Performance included AUROC, area under the precision–recall curve, Brier score, "
        "sensitivity, specificity, balanced accuracy, calibration intercept, and calibration "
        "slope. Decision curves used risk thresholds from 0.05 to 0.60 [17]."
    )
    doc.add_paragraph(
        "Compact panels used one stratified 75:25 split fixed by seed 42. Candidate ranking, "
        "forward selection of up to five markers, threshold selection, and regularization "
        "tuning occurred within the training partition. The test partition remained untouched "
        "until the panel was locked. Clinical, clinical-plus-APOE, and multimodal candidate "
        "pools were compared."
    )

    doc.add_heading("Statistical analysis", level=2)
    doc.add_paragraph(
        "Continuous baseline characteristics were compared with Welch tests. Binary measures "
        "used Fisher exact tests. Standardized mean differences described magnitude. AUROC "
        "confidence intervals used 3,000 stratified bootstrap samples. Adjacent AUROC changes "
        "used paired stratified bootstrap samples. Model discrimination against chance used "
        "the Mann–Whitney relation to AUROC. Continuous outcomes used ridge regression with "
        "five-fold cross-validation. Paired bootstrap samples tested changes in R². The "
        "Benjamini–Hochberg procedure controlled the false discovery rate within each family "
        "of comparisons [20]. Two-sided q<0.05 defined statistical significance. Analyses "
        "used Python 3.14, pandas, NumPy, SciPy, and scikit-learn. Reporting follows "
        "TRIPOD+AI and incorporates PROBAST+AI principles [18,19]."
    )

    doc.add_heading("External conceptual replication", level=2)
    doc.add_paragraph(
        "In OASIS-2, the index state was the first CDR 0.5 visit. Progression was CDR ≥1 "
        "within 24 months. Stable participants required 24 months of follow-up. Three "
        "cumulative sets were evaluated: demographics, MMSE, and derived MRI measures. "
        "Regularized logistic regression used leave-one-subject-out prediction. Bootstrap "
        "intervals used 5,000 stratified samples. Permutation tests compared AUROC with 0.5. "
        "This analysis tested whether the direction of the phenotyping-depth pattern appeared "
        "in a second dataset; it did not test the locked ADNI coefficients."
    )

    add_figure(
        doc, 1, "figure1_cohort_attrition.png",
        "ADNI cohort construction. Stable participants required 24 months without dementia. "
        "Progressors could meet the outcome before month 24.",
        width=6.25,
    )

    doc.add_heading("Results", level=1)
    doc.add_heading("Cohort", level=2)
    doc.add_paragraph(
        "The diagnosis files contained 3,777 participants with usable longitudinal diagnoses. "
        "Among 1,803 participants with at least one MCI visit, 624 had no later visit. A further "
        "175 had less than 24 months of follow-up without observed progression. The primary "
        "cohort contained 1,004 participants. Of these, 261 (26.0%) progressed to dementia "
        "within 24 months and 743 remained stable."
    )
    doc.add_paragraph(
        "Progressors had lower MMSE and higher CDR-SB at baseline. APOE ε4 carriage, amyloid "
        "PET positivity, and plasma p-tau181 also differed after false-discovery correction. "
        "Amyloid PET was available in 602 participants. Plasma p-tau181 was available in 429. "
        "These comparisons describe association and availability; they do not estimate "
        "incremental prediction."
    )

    baseline_rows = []
    for row in t1.itertuples(index=False):
        baseline_rows.append([
            row[0], row[1], row[2], row[3], str(int(row[4])),
            f"{row[5]:.2f}", fmt_q(row[7]),
        ])
    add_table(
        doc,
        "Table 1. Baseline characteristics by 24-month outcome",
        ["Characteristic", "Overall", "Stable MCI", "Progressor", "Available n", "SMD", "FDR q"],
        baseline_rows,
        [2050, 1280, 1280, 1280, 840, 680, 720],
        "Values are mean (SD) or n (%). SMD, standardized mean difference. "
        "FDR q values use Benjamini–Hochberg correction across the nine baseline comparisons.",
        font_size=7.8,
    )

    doc.add_heading("Stepwise classification", level=2)
    doc.add_paragraph(
        "Demographics alone did not discriminate progressors from stable participants "
        "(AUROC 0.510). The first change occurred when MMSE was added. MoCA, CDR-SB, and "
        "the broad cognitive battery each produced further gains after false-discovery "
        "correction. Broad cognition reached AUROC 0.866. MRI did not change AUROC. APOE and "
        "the biomarker layer produced small point increases, but neither increment was "
        "significant after correction."
    )
    step_rows = []
    for row in t3.itertuples(index=False):
        step_rows.append([
            str(int(row.step)),
            row.feature_set.split(" ", 1)[1] if " " in row.feature_set else row.feature_set,
            str(int(row.n_features)),
            fmt_ci(row.roc_auc, row.roc_auc_ci_low, row.roc_auc_ci_high),
            f"{row.pr_auc:.3f}",
            f"{row.brier_score:.3f}",
            "—" if pd.isna(row.delta_auc) else (
                "0.000" if abs(row.delta_auc) < 0.0005 else f"{row.delta_auc:+.3f}"),
            fmt_q(row.delta_auc_q),
        ])
    add_table(
        doc,
        "Table 2. Nested cross-validated classification performance",
        ["Step", "Cumulative feature set", "Features n", "AUROC (95% CI)", "PR-AUC", "Brier", "ΔAUROC", "FDR q"],
        step_rows,
        [540, 1870, 720, 1480, 760, 700, 800, 700],
        "AUROC intervals use 3,000 stratified bootstrap samples. ΔAUROC compares each row "
        "with the preceding cumulative set. PR-AUC, area under the precision–recall curve.",
        font_size=7.6,
    )

    add_figure(
        doc, 2, "figure2_stepwise_classification.png",
        "Discrimination across the cumulative phenotyping ladder. Error bars show 95% "
        "stratified-bootstrap confidence intervals. Asterisks mark adjacent increments with "
        "FDR q<0.05.",
    )
    add_figure(
        doc, 3, "figure3_performance_vs_burden.png",
        "AUROC plotted against a prespecified relative testing-burden score. The largest "
        "gain occurred before MRI, APOE, and biomarker testing.",
    )

    doc.add_heading("Calibration and decision utility", level=2)
    full = t3.iloc[-1]
    doc.add_paragraph(
        f"The full model’s mean predicted risk was {full['mean_predicted_risk']:.3f}, "
        f"compared with an observed event rate of {full['observed_event_rate']:.3f}. "
        f"The calibration intercept was {full['calibration_intercept']:.3f}. The calibration "
        f"slope was {full['calibration_slope']:.3f}. At a 0.50 threshold, sensitivity was "
        f"{full['sensitivity']:.3f} and specificity was {full['specificity']:.3f}. Decision "
        "curves showed positive net benefit over treat-none across part of the assessed "
        "threshold range. These estimates remain internally validated and should not be used "
        "for clinical decisions without transport validation."
    )
    add_figure(
        doc, 6, "figure6_discrimination_calibration_utility.png",
        "Out-of-fold discrimination, decile calibration, and decision-curve analysis for "
        "the broad-cognition model and the full biomarker model. The dashed diagonal in the "
        "calibration panel denotes agreement. Decision curves compare each model with "
        "treat-all and treat-none strategies.",
    )

    doc.add_heading("Compact panels", level=2)
    doc.add_paragraph(
        "The locked test set contained 251 participants, including 65 progressors. The "
        "five-marker clinical panel selected ADNI memory composite, FAQ total, ADNI executive "
        "composite, CDR-SB, and ADAS total. It reached AUROC 0.853. APOE was not selected in "
        "the clinical-plus-APOE pool, so that panel was identical. The multimodal panel "
        "replaced two clinical variables with right middle temporal thickness and left "
        "hippocampal volume. It did not outperform the clinical panel."
    )
    panel_rows = []
    for row in t5[t5.panel_size.eq(5)].itertuples(index=False):
        panel_rows.append([
            row.panel_family,
            row.features.replace(" | ", "; "),
            fmt_ci(row.roc_auc, row.roc_auc_ci_low, row.roc_auc_ci_high),
            f"{row.pr_auc:.3f}",
            f"{row.brier_score:.3f}",
            "—" if pd.isna(row.delta_auc_vs_clinical) else f"{row.delta_auc_vs_clinical:+.3f}",
            fmt_q(row.delta_auc_q),
        ])
    add_table(
        doc,
        "Table 3. Five-marker panels on the untouched test set",
        ["Panel", "Selected markers", "AUROC (95% CI)", "PR-AUC", "Brier", "Δ vs clinical", "FDR q"],
        panel_rows,
        [980, 2940, 1350, 700, 650, 850, 650],
        "Selection and tuning occurred only in the 75% training partition. Δ values compare "
        "paired test-set predictions with the five-marker clinical panel.",
        font_size=7.4,
    )
    add_figure(
        doc, 5, "figure5_minimal_panels.png",
        "Locked-test AUROC as markers were added within each candidate pool. The "
        "clinical-plus-APOE curve equals the clinical curve because APOE was not selected.",
    )

    doc.add_heading("Continuous outcomes", level=2)
    descriptions = {row.Outcome: row for row in t6.itertuples(index=False)}
    broad = t4[t4.feature_set.eq("6 + Broad cognition")].set_index("outcome")
    mri = t4[t4.feature_set.eq("7 + MRI")].set_index("outcome")
    biomarker = t4[t4.feature_set.eq("9 + AD biomarkers")].set_index("outcome")
    continuous_rows = []
    for outcome in ("MMSE change", "CDRSB change", "MOCA change"):
        label = outcome.replace("CDRSB", "CDR-SB").replace("MOCA", "MoCA")
        desc = descriptions[outcome]
        continuous_rows.append([
            label,
            str(int(desc.n)),
            f"{desc.Mean:+.2f} ({desc.SD:.2f})",
            f"{broad.loc[outcome, 'r2']:.3f}",
            f"{mri.loc[outcome, 'r2']:.3f}",
            fmt_q(mri.loc[outcome, "delta_r2_q"]),
            f"{biomarker.loc[outcome, 'r2']:.3f}",
            fmt_q(biomarker.loc[outcome, "delta_r2_q"]),
        ])
    add_table(
        doc,
        "Table 4. Prediction of approximately 24-month score change",
        ["Outcome", "n", "Mean change (SD)", "Broad cognition R²", "+MRI R²", "MRI Δ q", "Full R²", "Biomarker Δ q"],
        continuous_rows,
        [1080, 540, 1270, 1100, 850, 800, 800, 900],
        "Positive CDR-SB change denotes worsening. Negative MMSE or MoCA change denotes "
        "worsening. Δ q values compare MRI with broad cognition and the full model with MRI.",
        font_size=7.7,
    )
    doc.add_paragraph(
        "The broad cognitive set explained 33.9% of MMSE change variance and 29.1% of CDR-SB "
        "change variance. MRI did not add a significant R² increment. The full biomarker "
        "model also did not add a significant increment. MoCA change was available in 125 "
        "participants. Its broad-cognition increment was nominally positive but did not pass "
        "false-discovery correction."
    )
    add_figure(
        doc, 4, "figure4_continuous_outcomes.png",
        "Cross-validated R² for approximately 24-month change in CDR-SB, MMSE, and MoCA. "
        "Step labels match the cumulative phenotyping ladder.",
    )

    doc.add_heading("Sensitivity and subgroup analyses", level=2)
    sensitivity = t7.iloc[1]
    doc.add_paragraph(
        f"Relaxing the stable follow-up rule from 24 to 18 months increased the cohort to "
        f"{int(sensitivity['n']):,} participants. Full-model AUROC was "
        f"{sensitivity['ROC-AUC']:.3f} (95% confidence interval "
        f"{sensitivity['ROC-AUC CI low']:.3f}–{sensitivity['ROC-AUC CI high']:.3f}). "
        "No subgroup heterogeneity test passed FDR correction. The largest point difference "
        "was between women and men, but its corrected q value was "
        f"{t9.loc[(t9['Subgroup domain'].eq('Sex')) & (t9['Subgroup'].eq('Female')), 'Heterogeneity FDR q'].iloc[0]:.3f}."
    )

    doc.add_heading("OASIS-2 conceptual replication", level=2)
    oas_full = t8.iloc[-1]
    doc.add_paragraph(
        f"The OASIS-2 cohort contained {int(oas_full['n'])} eligible participants; "
        f"{int(oas_full['Converters n'])} progressed from CDR 0.5 to CDR ≥1 within 24 months. "
        f"AUROC increased from {t8.iloc[0]['ROC-AUC']:.3f} for demographics to "
        f"{t8.iloc[1]['ROC-AUC']:.3f} after MMSE and {oas_full['ROC-AUC']:.3f} after MRI. "
        f"The full interval was {oas_full['ROC-AUC CI low']:.3f}–"
        f"{oas_full['ROC-AUC CI high']:.3f}; q against chance was "
        f"{oas_full['ROC-AUC FDR q vs 0.5']:.3f}. Neither adjacent increment was significant. "
        "The point pattern was directionally compatible with added information, but the sample "
        "was too small for a precise test."
    )
    add_figure(
        doc, 7, "figure7_oasis2_external_replication.png",
        "Leave-one-subject-out AUROC in the OASIS-2 conceptual replication. Error bars show "
        "95% stratified-bootstrap confidence intervals. OASIS-2 used CDR-defined index and "
        "outcome states and a reduced feature space.",
    )

    doc.add_heading("Discussion", level=1)
    doc.add_heading("Principal findings", level=2)
    doc.add_paragraph(
        "Four findings answer the study question. First, demographic information did not "
        "separate short-horizon progressors from stable MCI. Second, MMSE, MoCA, CDR-SB, and "
        "the broader cognitive and functional battery produced the material gains. Third, "
        "MRI, APOE, and available AD biomarkers did not add significant discrimination after "
        "the clinical battery. Fourth, a five-marker clinical panel retained most of the "
        "performance on an untouched test set. The full model discriminated within ADNI, but "
        "the OASIS-2 analysis was too small and too different to establish transport."
    )

    doc.add_heading("Relation to prior work", level=2)
    doc.add_paragraph(
        "The result fits parts of the existing literature. A systematic review found cognition, "
        "MRI, APOE, age, and ADAS-Cog among recurrent predictors, while also noting limited "
        "external validation [11]. Da and colleagues reported that cognitive and MRI measures "
        "carried much of the prognostic information and that APOE and CSF biomarkers did not "
        "significantly improve their combined model [13]. Dickerson and Wolk found time-horizon "
        "differences: neurodegeneration markers were informative for short-term progression, "
        "while amyloid was more useful over a longer interval [12]. Our 24-month target and "
        "rich cognitive battery may therefore reduce the remaining increment available to "
        "biological measures."
    )
    doc.add_paragraph(
        "Van der Veere and colleagues reported limited explained variance for MMSE trajectories "
        "in amyloid-positive MCI and mild dementia, with modest transport estimates in ADNI "
        "[10]. Our continuous outcomes reached R² 0.34 for MMSE and 0.31 for CDR-SB at their "
        "best points, but MoCA estimates were unstable because follow-up coverage was small. "
        "Oxtoby and colleagues showed how progression staging can alter trial screening [9]. "
        "Our analysis complements that approach by quantifying measurement burden before a "
        "screening rule is chosen."
    )

    doc.add_heading("Biological interpretation", level=2)
    doc.add_paragraph(
        "The biomarker result requires a narrow interpretation. Amyloid positivity and plasma "
        "p-tau181 were associated with progression at baseline. APOE ε4 carriage was also more "
        "common among progressors. These findings align with the biological role of amyloid, "
        "tau, and APOE in AD [12–15]. Their later incremental AUROC changes were small because "
        "the comparison was conditional on extensive clinical phenotyping, used a short "
        "clinical outcome horizon, and relied on variables with uneven availability. A measure "
        "can identify AD pathology or inform long-term prognosis while adding little to "
        "short-term discrimination after current cognitive and functional status is known."
    )
    doc.add_paragraph(
        "The data do not support omitting biomarkers from diagnosis, treatment eligibility, or "
        "biological staging. The 2024 Alzheimer’s Association criteria give validated amyloid "
        "and phosphorylated tau biomarkers a diagnostic role [14]. APOE also affects disease "
        "risk and treatment discussions, with the strongest biological effects observed in "
        "ε4 homozygotes [15]. Our study addressed prediction of a clinical transition. It did "
        "not test diagnostic validity, treatment response, or causal mechanisms."
    )

    doc.add_heading("Implications", level=2)
    doc.add_paragraph(
        "For short-horizon risk stratification, a staged assessment is supported. Basic "
        "cognitive and functional testing should precede high-burden measurement. A compact "
        "clinical panel may provide a screening score for research recruitment. Biomarkers "
        "can then be reserved for questions they are designed to answer, including biological "
        "confirmation, treatment eligibility, or longer-horizon prognosis. This sequence "
        "requires prospective testing. The current model is not a clinical calculator."
    )

    doc.add_heading("Strengths and limitations", level=2)
    doc.add_paragraph(
        "Strengths include an explicit index state and prediction horizon, a stable-control "
        "follow-up rule, month-harmonized joins, fold-contained preprocessing, nested "
        "regularization, paired inference, multiple-testing correction, calibration, decision "
        "curves, a locked compact-panel test set, continuous outcomes, and reproducible code. "
        "The stepwise design answers a resource question that a single final model cannot."
    )
    doc.add_paragraph(
        "Several limitations affect interpretation. ADNI is an enriched research cohort and "
        "does not represent routine clinical populations. Diagnostic progression is not a "
        "biomarker-confirmed AD outcome. Missingness was substantial for MoCA and biological "
        "measures and was handled with median imputation plus missingness indicators; this "
        "cannot remove selection bias. Assays and imaging pipelines differed across ADNI "
        "phases. Plasma p-tau217 could not be analyzed because only eight baseline records were "
        "usable. The same ADNI cohort was used to compare phenotyping steps, so the results are "
        "internally rather than externally validated. The compact-panel test set was used once "
        "but still came from ADNI. OASIS-2 used CDR 0.5, a smaller feature set, and only 34 "
        "eligible participants. It is a conceptual replication. It does not satisfy direct "
        "external validation. No causal claim can be made from these observational data."
    )

    doc.add_heading("Conclusion", level=2)
    doc.add_paragraph(
        "In this ADNI cohort, cognitive and functional phenotyping supplied most of the "
        "information needed to predict dementia within 24 months after a first MCI visit. "
        "The full model achieved AUROC 0.875, but MRI, APOE, and available AD biomarkers did "
        "not add significant discrimination after the clinical battery. A five-marker clinical "
        "panel achieved locked-test AUROC 0.853. The next study should lock that panel, retain "
        "assay-specific biomarkers, and test it in a larger independent MCI cohort with "
        "biomarker-confirmed disease and prespecified follow-up."
    )

    doc.add_heading("Data and code availability", level=1)
    doc.add_paragraph(
        "Analysis code is available at https://github.com/kapoorsuvan-cpu/"
        "adni-phenotyping-depth-progression. Participant-level ADNI data and derived CSV files "
        "are not included. Approved investigators must obtain the ADNI exports through the "
        "LONI Image and Data Archive under the ADNI Data Use Agreement. The OASIS-2 subject "
        "workbook is obtained from Washington University under the OASIS data-use terms."
    )
    doc.add_paragraph(
        "To reproduce the study, place DXSUM.csv, PTDEMOG.csv, GDSCALE.csv, MMSE.csv, MOCA.csv, "
        "CDR.csv, FAQ.csv, ADAS.csv, NEUROBAT.csv, UWNPSYCHSUM.csv, UCSFFSX7.csv, APOERES.csv, "
        "UCBERKELEY_AMY_6MM.csv, UGOTPTAU181.csv, UPENNBIOMK_MASTER.csv, and "
        "UPENNBIOMK_ROCHE_ELECSYS.csv in the repository's `csvs/` directory. From the repository "
        "root, run `bash run_all.sh`. The entry point validates inputs with "
        "`scripts/check_inputs.py`, runs `scripts/run_final_analysis.py`, downloads and runs the "
        "OASIS-2 replication, and rebuilds the manuscript. `DATA_REQUIREMENTS.md` and "
        "`input_manifest.json` specify each source, filename, and required column."
    )

    doc.add_heading("Ethics", level=1)
    doc.add_paragraph(
        "This study analyzed de-identified secondary data. ADNI and OASIS investigators "
        "obtained approval and participant consent for the original studies. Researchers seeking "
        "to reproduce this analysis must obtain their own authorization to access ADNI and "
        "OASIS-2 data and comply with each repository's applicable data-use terms."
    )

    doc.add_heading("Author contributions", level=1)
    doc.add_paragraph(
        "Proposed Contributor Roles Taxonomy (CRediT) statement, subject to confirmation: Suvan "
        "Kapoor—Conceptualization, Methodology, Software, Formal Analysis, Visualization, Data "
        "Curation, Writing—Original Draft, and Project Administration. Dominic Ablakhad, Rayan "
        "Hanna, Kylan Huynh, Eric Quirarte, and Orion Nocon—Investigation, Validation, "
        "Interpretation, and Writing—Review and Editing."
    )

    doc.add_heading("Competing interests and funding", level=1)
    doc.add_paragraph(
        "Author competing-interest declarations and project-specific funding information must "
        "be confirmed before submission. No statement is inferred in this draft."
    )

    doc.add_heading("Acknowledgments", level=1)
    doc.add_paragraph(
        "The authors thank Douglas Galasko, MD, for methodological guidance."
    )
    doc.add_paragraph(
        "*Data used in preparation of this article were obtained from the Alzheimer’s Disease "
        "Neuroimaging Initiative (ADNI) database (adni.loni.usc.edu). As such, the investigators "
        "within the ADNI contributed to the design and implementation of ADNI and/or provided "
        "data but did not participate in analysis or writing of this report. A complete listing "
        "of ADNI investigators can be found at: "
        "http://adni.loni.usc.edu/wp-content/uploads/how_to_apply/ADNI_Acknowledgement_List.pdf"
    )
    doc.add_paragraph(
        "Data collection and sharing for the Alzheimer’s Disease Neuroimaging Initiative "
        "(ADNI) is funded by the National Institute on Aging (National Institutes of Health "
        "Grant U19 AG024904). The grantee organization is the Northern California Institute "
        "for Research and Education. In the past, ADNI has also received funding from the "
        "National Institute of Biomedical Imaging and Bioengineering, the Canadian Institutes "
        "of Health Research, and private-sector contributions through the Foundation for the "
        "National Institutes of Health."
    )
    doc.add_paragraph(
        "Data were provided in part by OASIS-2: Longitudinal MRI Data in Nondemented and "
        "Demented Older Adults. Principal investigators: D. Marcus, R. Buckner, J. Csernansky, "
        "and J. Morris. OASIS-2 support: P50 AG05681, P01 AG03991, P01 AG026276, R01 AG021910, "
        "P20 MH071616, and U24 RR021382."
    )

    doc.add_heading("References", level=1)
    references = [
        "1. Albert MS, DeKosky ST, Dickson D, et al. The diagnosis of mild cognitive impairment due to Alzheimer’s disease: recommendations from the National Institute on Aging–Alzheimer’s Association workgroups. Alzheimers Dement. 2011;7:270–279. doi:10.1016/j.jalz.2011.03.008.",
        "2. Petersen RC, Thomas RG, Grundman M, et al. Vitamin E and donepezil for the treatment of mild cognitive impairment. N Engl J Med. 2005;352:2379–2388. doi:10.1056/NEJMoa050151.",
        "3. Mueller SG, Weiner MW, Thal LJ, et al. The Alzheimer’s Disease Neuroimaging Initiative. Neuroimaging Clin N Am. 2005;15:869–877. doi:10.1016/j.nic.2005.09.008.",
        "4. Folstein MF, Folstein SE, McHugh PR. “Mini-mental state”: a practical method for grading the cognitive state of patients for the clinician. J Psychiatr Res. 1975;12:189–198. doi:10.1016/0022-3956(75)90026-6.",
        "5. Nasreddine ZS, Phillips NA, Bédirian V, et al. The Montreal Cognitive Assessment, MoCA: a brief screening tool for mild cognitive impairment. J Am Geriatr Soc. 2005;53:695–699. doi:10.1111/j.1532-5415.2005.53221.x.",
        "6. Morris JC. The Clinical Dementia Rating: current version and scoring rules. Neurology. 1993;43:2412–2414. doi:10.1212/WNL.43.11.2412-a.",
        "7. Crane PK, Carle A, Gibbons LE, et al. Development and assessment of a composite score for memory in the Alzheimer’s Disease Neuroimaging Initiative. Brain Imaging Behav. 2012;6:502–516. doi:10.1007/s11682-012-9186-z.",
        "8. Gibbons LE, Carle AC, Mackin RS, et al. A composite score for executive functioning, validated in ADNI participants with baseline mild cognitive impairment. Brain Imaging Behav. 2012;6:517–527. doi:10.1007/s11682-012-9176-1.",
        "9. Oxtoby NP, Shand C, Cash DM, Alexander DC, Barkhof F. Targeted screening for Alzheimer’s disease clinical trials using data-driven disease progression models. Front Artif Intell. 2022;5:660581. doi:10.3389/frai.2022.660581.",
        "10. van der Veere PJ, Hoogland J, Visser LNC, et al. Predicting cognitive decline in amyloid-positive patients with mild cognitive impairment or mild dementia. Neurology. 2024;103:e209605. doi:10.1212/WNL.0000000000209605.",
        "11. Chen Y, Qian X, Zhang Y, et al. Prediction models for conversion from mild cognitive impairment to Alzheimer’s disease: a systematic review and meta-analysis. Front Aging Neurosci. 2022;14:840386. doi:10.3389/fnagi.2022.840386.",
        "12. Dickerson BC, Wolk DA. Biomarker-based prediction of progression in MCI: comparison of AD signature and hippocampal volume with spinal fluid amyloid-β and tau. Front Aging Neurosci. 2013;5:55. doi:10.3389/fnagi.2013.00055.",
        "13. Da X, Toledo JB, Zee J, et al. Integration and relative value of biomarkers for prediction of MCI to AD progression: spatial patterns of brain atrophy, cognitive scores, APOE genotype and CSF biomarkers. Neuroimage Clin. 2014;4:164–173. doi:10.1016/j.nicl.2013.11.010.",
        "14. Jack CR Jr, Andrews JS, Beach TG, et al. Revised criteria for diagnosis and staging of Alzheimer’s disease: Alzheimer’s Association Workgroup. Alzheimers Dement. 2024;20:5143–5169. doi:10.1002/alz.13859.",
        "15. Fortea J, Pegueroles J, Alcolea D, et al. APOE4 homozygosity represents a distinct genetic form of Alzheimer’s disease. Nat Med. 2024;30:1284–1291. doi:10.1038/s41591-024-02931-w.",
        "16. Marcus DS, Fotenos AF, Csernansky JG, Morris JC, Buckner RL. Open Access Series of Imaging Studies: longitudinal MRI data in nondemented and demented older adults. J Cogn Neurosci. 2010;22:2677–2684. doi:10.1162/jocn.2009.21407.",
        "17. Vickers AJ, Elkin EB. Decision curve analysis: a novel method for evaluating prediction models. Med Decis Making. 2006;26:565–574. doi:10.1177/0272989X06295361.",
        "18. Collins GS, Moons KGM, Dhiman P, et al. TRIPOD+AI statement: updated guidance for reporting clinical prediction models that use regression or machine learning methods. BMJ. 2024;385:e078378. doi:10.1136/bmj-2023-078378.",
        "19. Moons KGM, Damen JAA, Kaul T, et al. PROBAST+AI: an updated quality, risk of bias, and applicability assessment tool for prediction models using regression or artificial intelligence methods. BMJ. 2025;388:e082505. doi:10.1136/bmj-2024-082505.",
        "20. Benjamini Y, Hochberg Y. Controlling the false discovery rate: a practical and powerful approach to multiple testing. J R Stat Soc B. 1995;57:289–300.",
    ]
    for reference in references:
        p = doc.add_paragraph(reference)
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(3)
        for run in p.runs:
            run.font.size = Pt(8.8)

    PUBLICATION.mkdir(parents=True, exist_ok=True)
    doc.save(MANUSCRIPT)


def build_supplement() -> None:
    doc = Document()
    configure_document(doc, "Supplement: phenotyping depth in MCI progression")
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Supplementary Material")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        "How Much Phenotyping Is Needed to Predict 24-Month Progression "
        "From Mild Cognitive Impairment to Dementia?"
    ).italic = True
    authors = doc.add_paragraph(
        "Suvan Kapoor; Dominic Ablakhad; Rayan Hanna; Kylan Huynh; "
        "Eric Quirarte; Orion Nocon"
    )
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_rule(doc)
    doc.add_paragraph(
        "This supplement is generated from the same CSV outputs as the manuscript. "
        "Values are not transcribed manually."
    )

    table_specs = [
        (
            "Supplementary Table S1. Feature coverage",
            "table2_feature_coverage.csv",
            ["Feature domain", "Features n", "Participants with any feature n", "Participants with any feature %", "Median feature availability %"],
        ),
        (
            "Supplementary Table S2. Full stepwise classification output",
            "table3_stepwise_classification.csv",
            ["step", "feature_set", "n_features", "roc_auc", "roc_auc_ci_low", "roc_auc_ci_high", "pr_auc", "brier_score", "delta_auc", "delta_auc_ci_low", "delta_auc_ci_high", "delta_auc_p", "delta_auc_q"],
        ),
        (
            "Supplementary Table S3. Full continuous-outcome output",
            "table4_continuous_outcomes.csv",
            ["outcome", "feature_set", "n", "n_features", "r2", "rmse", "mae", "delta_r2", "delta_r2_ci_low", "delta_r2_ci_high", "delta_r2_p", "delta_r2_q"],
        ),
        (
            "Supplementary Table S4. Compact-panel sequence",
            "table5_minimal_marker_panels.csv",
            ["panel_family", "panel_size", "features", "selected_C", "test_n", "roc_auc", "roc_auc_ci_low", "roc_auc_ci_high", "pr_auc", "brier_score", "delta_auc_vs_clinical", "delta_auc_p", "delta_auc_q"],
        ),
        (
            "Supplementary Table S5. Outcome descriptives",
            "table6_outcome_descriptives.csv",
            ["Outcome", "n", "Mean", "SD", "Median", "Q1", "Q3"],
        ),
        (
            "Supplementary Table S6. Stable-follow-up sensitivity analysis",
            "table7_sensitivity_cohorts.csv",
            ["Stable follow-up rule", "n", "Converters n", "Stable n", "ROC-AUC", "ROC-AUC CI low", "ROC-AUC CI high", "PR-AUC", "Brier score"],
        ),
        (
            "Supplementary Table S7. OASIS-2 conceptual replication",
            "table8_oasis2_external_replication.csv",
            ["Feature set", "n", "Converters n", "Stable n", "ROC-AUC", "ROC-AUC CI low", "ROC-AUC CI high", "PR-AUC", "Brier score", "ROC-AUC p vs 0.5", "ROC-AUC FDR q vs 0.5", "Delta AUC", "Delta AUC CI low", "Delta AUC CI high", "Delta AUC p", "Delta AUC FDR q"],
        ),
        (
            "Supplementary Table S8. Subgroup performance",
            "table9_subgroup_performance.csv",
            ["Subgroup domain", "Subgroup", "n", "Converters n", "ROC-AUC", "ROC-AUC CI low", "ROC-AUC CI high", "PR-AUC", "Brier score", "Second-minus-first AUC difference", "Difference CI low", "Difference CI high", "Heterogeneity p", "Heterogeneity FDR q"],
        ),
    ]

    for title_text, filename, columns in table_specs:
        data = load_table(filename)[columns].copy()
        formatted_rows = []
        for row in data.itertuples(index=False, name=None):
            values = []
            for value in row:
                if pd.isna(value):
                    values.append("—")
                elif isinstance(value, (float, np.floating)):
                    if abs(value) < 0.001 and value != 0:
                        values.append(f"{value:.2e}")
                    else:
                        values.append(f"{value:.3f}")
                else:
                    values.append(str(value))
            formatted_rows.append(values)
        widths = [max(500, int(9300 / len(columns))) for _ in columns]
        if len(columns) > 10:
            section = doc.add_section(WD_SECTION.NEW_PAGE)
            section.orientation = 1
            section.page_width, section.page_height = section.page_height, section.page_width
            section.top_margin = Inches(0.55)
            section.bottom_margin = Inches(0.55)
            section.left_margin = Inches(0.55)
            section.right_margin = Inches(0.55)
            total = 13900
            widths = [max(650, int(total / len(columns))) for _ in columns]
        add_table(
            doc, title_text, [str(c).replace("_", " ") for c in columns],
            formatted_rows, widths,
            "Full machine-readable values are retained in the corresponding CSV file.",
            font_size=6.4 if len(columns) > 10 else 7.2,
        )
        doc.add_page_break()

    doc.add_heading("Supplementary reproducibility record", level=1)
    doc.add_paragraph(
        "Random seed: 42. Primary horizon: 24 months. Stable follow-up requirement: "
        "24 months. Outer validation: five-fold stratified out-of-fold prediction. Inner "
        "tuning: four-fold grid search by negative log loss. Candidate C values: 0.003, "
        "0.01, 0.03, 0.1, 0.3, and 1. Bootstrap repetitions: 3,000 for ADNI AUROC and "
        "increment tests; 5,000 for OASIS-2 AUROC intervals and increment tests."
    )
    doc.add_paragraph(
        "Machine-readable supporting files include the analysis cohort, out-of-fold "
        "probabilities, join audit, patient-level continuous outcomes, compact-marker "
        "recurrence, analysis manifest, and OASIS-2 provenance checksum."
    )

    PUBLICATION.mkdir(parents=True, exist_ok=True)
    doc.save(SUPPLEMENT)


def main() -> None:
    build_main()
    build_supplement()
    print(f"Created: {MANUSCRIPT}")
    print(f"Created: {SUPPLEMENT}")


if __name__ == "__main__":
    main()
